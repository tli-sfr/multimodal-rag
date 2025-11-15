"""Text document ingestion with semantic chunking."""

from pathlib import Path
from typing import List, Optional
from uuid import uuid4

from langchain_text_splitters import RecursiveCharacterTextSplitter
from loguru import logger
import pypdf
from docx import Document as DocxDocument

from ..models import Document, Chunk, ModalityType
from ..config import get_config
from .base import BaseIngester


class TextIngester(BaseIngester):
    """Ingester for text documents (PDF, TXT, DOCX)."""
    
    SUPPORTED_EXTENSIONS = {'.txt', '.pdf', '.docx', '.md'}
    
    def __init__(
        self,
        chunk_size: Optional[int] = None,
        chunk_overlap: Optional[int] = None,
    ):
        """Initialize text ingester.
        
        Args:
            chunk_size: Size of text chunks
            chunk_overlap: Overlap between chunks
        """
        super().__init__(ModalityType.TEXT)
        
        config = get_config()
        self.chunk_size = chunk_size or config.get('ingestion.chunking.chunk_size', 512)
        self.chunk_overlap = chunk_overlap or config.get('ingestion.chunking.chunk_overlap', 50)
        
        self.text_splitter = RecursiveCharacterTextSplitter(
            chunk_size=self.chunk_size,
            chunk_overlap=self.chunk_overlap,
            length_function=len,
            separators=["\n\n", "\n", ". ", " ", ""]
        )
    
    def validate_file(self, file_path: Path) -> bool:
        """Validate text file.
        
        Args:
            file_path: Path to file
            
        Returns:
            True if valid
        """
        if not self._validate_file_exists(file_path):
            return False
        
        if file_path.suffix.lower() not in self.SUPPORTED_EXTENSIONS:
            logger.error(f"Unsupported file extension: {file_path.suffix}")
            return False
        
        if not self._validate_file_size(file_path):
            return False
        
        return True
    
    def ingest(self, file_path: Path, **kwargs) -> Document:
        """Ingest text document.
        
        Args:
            file_path: Path to file
            **kwargs: Additional parameters
            
        Returns:
            Document object with chunks
        """
        if not self.validate_file(file_path):
            raise ValueError(f"Invalid file: {file_path}")
        
        logger.info(f"Ingesting text document: {file_path}")
        
        # Extract text based on file type
        text = self._extract_text(file_path)
        
        if not text or len(text.strip()) == 0:
            raise ValueError(f"No text extracted from {file_path}")
        
        # Create metadata
        metadata = self.create_metadata(file_path)
        
        # Create document
        doc_id = uuid4()
        document = Document(
            id=doc_id,
            title=file_path.stem,
            content=text,
            modality=ModalityType.TEXT,
            metadata=metadata
        )
        
        # Create chunks
        chunks = self._create_chunks(text, doc_id, metadata)
        document.chunks = chunks
        
        logger.info(f"Created {len(chunks)} chunks from {file_path}")
        
        return document
    
    def _extract_text(self, file_path: Path) -> str:
        """Extract text from file based on extension.
        
        Args:
            file_path: Path to file
            
        Returns:
            Extracted text
        """
        suffix = file_path.suffix.lower()
        
        try:
            if suffix == '.txt' or suffix == '.md':
                return self._extract_from_txt(file_path)
            elif suffix == '.pdf':
                return self._extract_from_pdf(file_path)
            elif suffix == '.docx':
                return self._extract_from_docx(file_path)
            else:
                raise ValueError(f"Unsupported file type: {suffix}")
        
        except Exception as e:
            logger.error(f"Error extracting text from {file_path}: {e}")
            raise
    
    def _extract_from_txt(self, file_path: Path) -> str:
        """Extract text from TXT file."""
        with open(file_path, 'r', encoding='utf-8', errors='ignore') as f:
            return f.read()
    
    def _extract_from_pdf(self, file_path: Path) -> str:
        """Extract text from PDF file."""
        text_parts = []
        
        with open(file_path, 'rb') as f:
            pdf_reader = pypdf.PdfReader(f)
            
            for page_num, page in enumerate(pdf_reader.pages):
                page_text = page.extract_text()
                if page_text:
                    text_parts.append(page_text)
        
        return "\n\n".join(text_parts)
    
    def _extract_from_docx(self, file_path: Path) -> str:
        """Extract text from DOCX file."""
        doc = DocxDocument(file_path)
        text_parts = [paragraph.text for paragraph in doc.paragraphs if paragraph.text]
        return "\n\n".join(text_parts)
    
    def _create_chunks(
        self,
        text: str,
        parent_id: uuid4,
        metadata
    ) -> List[Chunk]:
        """Create chunks from text.
        
        Args:
            text: Full text
            parent_id: Parent document ID
            metadata: Document metadata
            
        Returns:
            List of chunks
        """
        # Split text into chunks
        text_chunks = self.text_splitter.split_text(text)
        
        # Create Chunk objects
        chunks = []
        for idx, chunk_text in enumerate(text_chunks):
            chunk = Chunk(
                content=chunk_text,
                modality=ModalityType.TEXT,
                metadata=metadata,
                chunk_index=idx,
                parent_id=parent_id
            )
            chunks.append(chunk)
        
        return chunks

